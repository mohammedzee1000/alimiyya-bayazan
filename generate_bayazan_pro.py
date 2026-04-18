import os
import sqlite3
import argparse
import shutil
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from config import PATHS, COLORS, STYLE, POS_MAP, THEMES, DEFAULT_THEME, get_theme_config, list_available_themes

def to_arabic_numerals(num):
    """Convert English numerals to Arabic-Indic numerals."""
    arabic_digits = '٠١٢٣٤٥٦٧٨٩'
    return ''.join(arabic_digits[int(d)] for d in str(num))

class BayazanProEngine:
    def __init__(self, theme=None):
        """
        Initialize the Pro Academic Engine with theme support.
        
        Args:
            theme: Theme name ('indopak' or 'uthmani'). If None, uses default theme.
        """
        self.theme_name = theme or DEFAULT_THEME
        self.theme_config = get_theme_config(self.theme_name)
        
        # Validate theme resources exist
        if not os.path.exists(self.theme_config["text_db"]):
            raise FileNotFoundError(
                f"Theme '{self.theme_name}' text database not found: {self.theme_config['text_db']}\n"
                f"Please ensure the theme resources are properly installed."
            )
        
        self.metadata = self._parse_metadata()
        self.morph_map = self._parse_leeds_morphology()
        # Para/Juz mapping for Juz 30
        self.juz_map = {i: 30 for i in range(78, 115)} 
        
        print(f"🎨 Using theme: {self.theme_config['name']}")
        print(f"📝 Font: {self.theme_config['font_name']}")

    def _parse_metadata(self):
        tree = ET.parse(PATHS["metadata_xml"])
        root = tree.getroot()
        data = {}
        for s in root.find('suras'):
            idx = int(s.get('index'))
            data[idx] = {
                "name": s.get('name'),
                "arname": s.get('arname') or s.get('name'),
                "ayas": int(s.get('ayas'))
            }
        return data

    def _parse_leeds_morphology(self):
        mapping = {}
        leeds_path = PATHS["morphology"]["leeds_txt"]
        if not os.path.exists(leeds_path): return mapping
        with open(leeds_path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.startswith('('):
                    parts = line.split('\t')
                    loc_raw = parts[0].strip('()').split(':')
                    loc = f"{loc_raw[0]}:{loc_raw[1]}:{loc_raw[2]}"
                    pos = parts[2]
                    if loc not in mapping or mapping[loc] == 'PARTICLE':
                        mapping[loc] = POS_MAP.get(pos, 'PARTICLE')
        return mapping

    def is_structural_symbol(self, text):
        """Check if text is a structural symbol (sajdah, ruku, waqf marks, etc.)"""
        if len(text) == 1 and ord(text[0]) > 0x06FF: return True
        
        # Comprehensive list of Quranic structural markers
        markers = [
            # Sajdah markers
            '۩',
            # Ruku markers
            '۞',
            # Waqf (pause) marks
            'ۖ', 'ۗ', 'ۘ', 'ۙ', 'ۚ', 'ۛ', 'ۜ',
            # Small high signs
            'ۣ', 'ۤ', 'ۥ', 'ۦ',
            # Maddah and other marks
            '۟', '۠',
            # Quranic annotation marks
            '۝', '۞',
            # Additional pause marks
            'ۗ', 'ۖ'
        ]
        return any(m in text for m in markers)

    def set_cell_shading(self, cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)

    def set_arabic_font(self, run, size, color=None):
        """Set Arabic font using the current theme's font configuration."""
        font_name = self.theme_config["font_name"]
        run.font.name = font_name
        run.font.size = Pt(size)
        if color: run.font.color.rgb = RGBColor(*color)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)

    def get_enriched_words(self, sura, ayah):
        """Get words with morphological enrichment using theme's text database."""
        conn_text = sqlite3.connect(self.theme_config["text_db"])
        conn_root = sqlite3.connect(PATHS["morphology"]["root"])
        words = conn_text.execute(
            "SELECT location, text FROM words WHERE surah=? AND ayah=? ORDER BY word",
            (sura, ayah)
        ).fetchall()
        
        clean_header = " ".join([w[1] for w in words if not self.is_structural_symbol(w[1])])
        # Store both the text and the ayah number separately for proper font handling
        arabic_num = to_arabic_numerals(ayah)
        full_ref = f"{clean_header} \u200f﴿{arabic_num}﴾\u200f"
        
        enriched = []
        for loc, text in words:
            is_sym = self.is_structural_symbol(text)
            root_formatted = ""
            if not is_sym:
                root_res = conn_root.execute(
                    "SELECT r.arabic_trilateral FROM roots r JOIN root_words rw ON r.id = rw.root_id WHERE rw.word_location = ?", 
                    (loc,)
                ).fetchone()
                if root_res: root_formatted = root_res[0].strip().replace(" ", " . ")
            
            color = (166, 166, 166) if is_sym else COLORS.get(self.morph_map.get(loc, 'DEFAULT'), (0,0,0))
            enriched.append({"text": text, "root": root_formatted, "color": color, "is_symbol": is_sym})
        conn_text.close(); conn_root.close()
        return full_ref, enriched

    def add_page_number(self, paragraph):
        """Injects dynamic Word PAGE field into a paragraph."""
        run = paragraph.add_run()
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), 'PAGE')
        run._r.append(fldSimple)

    def add_table_of_contents(self, paragraph):
        """Injects a native MS Word Table of Contents (TOC) field."""
        run1 = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run1._r.append(fldChar1)

        run2 = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-1" \\h \\z \\u'
        run2._r.append(instrText)

        run3 = paragraph.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run3._r.append(fldChar2)

        run4 = paragraph.add_run(" Right-click here and select 'Update Field' to generate the Index.")
        run4.font.italic = True
        run4.font.color.rgb = RGBColor(150, 150, 150)

        run5 = paragraph.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run5._r.append(fldChar3)

    def create_workbook(self, start, end, output_filename):
        doc = Document()
        
        # Set core document properties to ensure modern format
        doc.core_properties.title = f"Alimiyya Bayazan - Surah {start} to {end}"
        doc.core_properties.author = "Alimiyya Bayazan Generator"
        doc.core_properties.comments = f"Generated with theme: {self.theme_config['name']}"
        
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Inches(11.69), Inches(8.27)
        section.left_margin = section.right_margin = Inches(0.5)

        # Add footer with page numbers
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_page_number(footer_para)

        # 1. Cover Page
        doc.add_paragraph("\n\n")
        
        p_taawwudh = doc.add_paragraph()
        p_taawwudh.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_taawwudh = p_taawwudh.add_run("أَعُوذُ بِاللَّهِ مِنَ الشَّيْطَانِ الرَّجِيمِ")
        self.set_arabic_font(run_taawwudh, 28)

        p_bismillah = doc.add_paragraph()
        p_bismillah.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_bismillah = p_bismillah.add_run("بِسْمِ اللَّهِ الرَّحْمَـٰنِ الرَّحِيمِ")
        self.set_arabic_font(run_bismillah, 28)
        
        doc.add_paragraph("\n")
        
        # Title Block
        p_title_block = doc.add_paragraph()
        p_title_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_main_title = p_title_block.add_run("Alimiyya Bayazan Pro Workbook\n")
        run_main_title.bold = True
        run_main_title.font.size = Pt(28)
        run_main_title.font.color.rgb = RGBColor(0, 51, 102)
        
        run_sub_title = p_title_block.add_run(f"Covering Surah {start} to Surah {end}\n")
        run_sub_title.font.size = Pt(18)
        run_sub_title.font.color.rgb = RGBColor(100, 100, 100)
        
        run_theme = p_title_block.add_run(f"Theme: {self.theme_config['name']}\n\n")
        run_theme.font.size = Pt(14)
        run_theme.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph("\n")
        
        # 2. Table of Contents
        p_index_title = doc.add_paragraph()
        p_index_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_index_title = p_index_title.add_run("Index / الفهرس")
        run_index_title.bold = True
        run_index_title.font.size = Pt(16)

        p_note = doc.add_paragraph()
        p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_note = p_note.add_run("📌 Quick Note: To calculate dynamic page numbers, right-click the grey text below and select 'Update Field'.")
        run_note.font.italic = True
        run_note.font.size = Pt(10)
        run_note.font.color.rgb = RGBColor(128, 128, 128)
        
        p_toc = doc.add_paragraph()
        p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_table_of_contents(p_toc)
        
        doc.add_page_break()

        for s_num in range(start, end + 1):
            s_info = self.metadata[s_num]
            juz_num = self.juz_map.get(s_num, "??")
            
            # 2. BIG Surah Header with Juz Info
            doc.add_paragraph(f"Juz {juz_num}", style='Caption').alignment = WD_ALIGN_PARAGRAPH.RIGHT
            h_ar = doc.add_heading('', level=1)
            h_ar.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            h_ar.paragraph_format.rtl = True
            run_h = h_ar.add_run(f"سورة {s_info['arname']}")
            self.set_arabic_font(run_h, 36, color=(0, 51, 102))

            # Surah-level notes section
            p_intro = doc.add_paragraph("Surah Introduction & Key Themes:")
            p_intro.runs[0].bold = True
            p_intro.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            doc.add_paragraph("\n\n\n")  # Space for notes

            if s_num != 9 and s_num != 1:
                p_bis = doc.add_paragraph()
                p_bis.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_bis.paragraph_format.rtl = True
                run_bis = p_bis.add_run("بِسْمِ اللَّهِ الرَّحْمَـٰنِ الرَّحِيمِ")
                self.set_arabic_font(run_bis, 22)

            for a_idx in range(1, s_info['ayas'] + 1):
                full_ref, words = self.get_enriched_words(s_num, a_idx)
                
                # 3. English "Ayah X:" label on separate line
                p_ayah_num = doc.add_paragraph()
                p_ayah_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run_ayah_label = p_ayah_num.add_run(f"Ayah {a_idx}:")
                run_ayah_label.font.size = Pt(14)
                run_ayah_label.font.color.rgb = RGBColor(100, 100, 100)
                run_ayah_label.bold = True
                
                # 4. Ayah Text with Arabic-Indic number at the end
                p_full = doc.add_paragraph()
                p_full.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p_full.paragraph_format.rtl = True
                
                # Split the text and number for different font handling
                # Extract text without the ayah number
                ayah_text_only = full_ref.rsplit('﴿', 1)[0].strip()
                arabic_num = to_arabic_numerals(a_idx)
                
                # Add ayah text with theme font
                run_text = p_full.add_run(ayah_text_only + " ")
                self.set_arabic_font(run_text, 20, color=(96, 96, 96))
                
                # Add ayah number with square brackets and Arabic-Indic numerals
                # Using square brackets for reliable rendering across all fonts
                run_num = p_full.add_run(f"[{arabic_num}]")
                run_num.font.name = "Arial"
                run_num.font.size = Pt(16)
                run_num.font.bold = True
                run_num.font.color.rgb = RGBColor(96, 96, 96)
                # Set complex script font
                r = run_num._element
                rPr = r.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:cs'), 'Arial')
                rPr.append(rFonts)

                # 5. Academic Analysis Table
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                table.allow_autofit = False
                
                table.columns[0].width = Inches(1.2)
                table.columns[1].width = Inches(1.1)
                table.columns[2].width = Inches(2.3)
                table.columns[3].width = Inches(5.9) # Notes area

                hdr_cells = table.rows[0].cells
                # Updated Header List
                headers = ["Word", "Root", "Sarf", "Notes / ملاحظات"]
                
                # Fallback logic for background color to prevent crashes
                bg_color = STYLE.get("table_header_bg", STYLE.get("header_bg", "F2F2F2"))
                
                for i, txt in enumerate(headers):
                    hdr_cells[i].text = txt
                    self.set_cell_shading(hdr_cells[i], bg_color)
                    hdr_cells[i].paragraphs[0].runs[0].bold = True

                for w in words:
                    row = table.add_row().cells
                    p_w = row[0].paragraphs[0]
                    p_w.paragraph_format.rtl = True
                    # Vertical breathing room for handwritten notes
                    p_w.paragraph_format.line_spacing = 1.2 
                    p_w.alignment = WD_ALIGN_PARAGRAPH.CENTER if w['is_symbol'] else WD_ALIGN_PARAGRAPH.RIGHT
                    
                    run_w = p_w.add_run(w['text'])
                    self.set_arabic_font(run_w, STYLE["font_size_arabic"], color=w['color'])
                    
                    p_r = row[1].paragraphs[0]
                    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_r = p_r.add_run(w['root'])
                    self.set_arabic_font(run_r, 14)

                doc.add_paragraph("\nAyah Context / General Analysis:")
                doc.add_paragraph("________________________________________________________________________________")

            doc.add_page_break()

        # Ensure output directory exists
        os.makedirs(PATHS["output_dir"], exist_ok=True)
        
        output_path = os.path.join(PATHS["output_dir"], output_filename)
        doc.save(output_path)
        
        # Copy font file to generated directory
        self._copy_font_to_output()
        
        # Create README if it doesn't exist
        self._create_font_readme()
        
        print(f"✅ Success! Workbook generated: {output_filename}")

    def _copy_font_to_output(self):
        """Copy the theme's font file to the generated directory."""
        font_source = self.theme_config["font_file"]
        if os.path.exists(font_source):
            font_filename = os.path.basename(font_source)
            font_dest = os.path.join(PATHS["output_dir"], font_filename)
            
            # Only copy if it doesn't exist or is different
            if not os.path.exists(font_dest):
                shutil.copy2(font_source, font_dest)
                print(f"📝 Font copied: {font_filename}")
        else:
            print(f"⚠️  Warning: Font file not found: {font_source}")

    def _create_font_readme(self):
        """Create a README file in the generated directory with font installation instructions."""
        readme_path = os.path.join(PATHS["output_dir"], "FONT_INSTALLATION.md")
        
        # Only create if it doesn't exist
        if not os.path.exists(readme_path):
            readme_content = f"""# Font Installation Required

## ⚠️ Important: Install Font Before Opening Documents

The generated Word documents in this directory use a specific Arabic font for proper text rendering. You **must** install the font before opening the documents to ensure correct display.

### Current Theme: {self.theme_config['name']}
**Font Required:** `{self.theme_config['font_name']}`

### Installation Instructions

#### Windows:
1. Locate the font file: `{os.path.basename(self.theme_config['font_file'])}`
2. Right-click on the font file
3. Select "Install" or "Install for all users"
4. Restart Microsoft Word if it's already open

#### macOS:
1. Double-click the font file: `{os.path.basename(self.theme_config['font_file'])}`
2. Click "Install Font" in Font Book
3. Restart Microsoft Word if it's already open

#### Linux:
1. Copy the font file to `~/.fonts/` or `/usr/share/fonts/`
2. Run: `fc-cache -f -v`
3. Restart your document viewer

### Verification

After installing the font:
1. Open any generated `.docx` file
2. The Arabic text should display clearly with proper diacritics
3. If text appears as boxes or incorrect characters, the font is not properly installed

### Font Sources

- **Indo-Pak Theme**: AlQuran IndoPak by QuranWBW
- **Uthmani Theme**: (To be added)
- **Standard Mode**: KFGQPC Nastaleeq Regular from [King Fahd Complex](https://fonts.qurancomplex.gov.sa/)

### Troubleshooting

If the text still doesn't display correctly after installation:
1. Ensure you installed the correct font file
2. Restart your computer (not just Word)
3. Check that the font appears in your system's font list
4. Try opening the document in a different Word version

---

**Generated by Alimiyya Bayazan Generator**
"""
            
            with open(readme_path, 'w', encoding='utf-8') as f:
                f.write(readme_content)
            print(f"📄 Created: FONT_INSTALLATION.md")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate Pro Academic Quranic Workbooks with theme support",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=f"""
Available Themes:
{chr(10).join(f"  • {name}: {desc}" for name, desc in list_available_themes().items())}

Examples:
  python generate_bayazan_pro.py --start 78 --end 114 --theme indopak
  python generate_bayazan_pro.py --start 1 --end 2 --theme uthmani -o "My_Workbook.docx"
        """
    )
    parser.add_argument("--start", type=int, required=True, help="Starting Surah number (1-114)")
    parser.add_argument("--end", type=int, required=True, help="Ending Surah number (1-114)")
    parser.add_argument("-o", "--output", type=str, help="Custom output filename")
    parser.add_argument("--theme", type=str, default=DEFAULT_THEME, 
                       choices=list(THEMES.keys()),
                       help=f"Theme to use (default: {DEFAULT_THEME})")
    args = parser.parse_args()
    
    filename = args.output or f"Academic_Bayazan_{args.theme}_{args.start}_{args.end}.docx"
    
    try:
        engine = BayazanProEngine(theme=args.theme)
        engine.create_workbook(args.start, args.end, filename)
    except FileNotFoundError as e:
        print(f"❌ Error: {e}")
        exit(1)
    except ValueError as e:
        print(f"❌ Error: {e}")
        exit(1)

# Made with Bob
