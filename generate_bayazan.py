import os
import shutil
import docx
import argparse
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from config import PATHS

# --- CONFIGURATION ---
INPUT_FILE = PATHS["standard"]["text_file"]
ARABIC_FONT = PATHS["standard"]["font_name"]

def set_arabic_font(run, size):
    """Forces macOS Word to respect the complex script Arabic font."""
    run.font.name = ARABIC_FONT
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:cs'), ARABIC_FONT)
    rPr.append(rFonts)

def set_cell_background(cell, hex_color):
    """Injects raw XML to apply a background shade to a Word table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:shd")
    if tcBorders is not None:
        tcPr.remove(tcBorders)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    tcPr.append(shd)

def add_page_number(paragraph):
    """Injects dynamic Word PAGE field into a paragraph."""
    run = paragraph.add_run()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    run._r.append(fldSimple)

def add_table_of_contents(paragraph):
    """Injects a native MS Word Table of Contents (TOC) field."""
    run1 = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-1" \\h \\z \\u' 
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run(" Right-click here and select 'Update Field' to generate the Index.")
    run4.font.italic = True
    run4.font.color.rgb = RGBColor(150, 150, 150)

    run5 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)

def create_doc(start_surah, end_surah, output_filepath):
    doc = docx.Document()
    
    # --- PAGE SETUP ---
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # --- FOOTER (PAGE NUMBERS) ---
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_para)

    print(f"Compiling Alimiyya Bayazan matrix into '{output_filepath}'...")

    # --- COVER PAGE: TITLE & BLESSINGS ---
    doc.add_paragraph("\n\n")

    p_taawwudh = doc.add_paragraph()
    p_taawwudh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_taawwudh = p_taawwudh.add_run("أَعُوذُ بِاللَّهِ مِنَ الشَّيْطَانِ الرَّجِيمِ")
    set_arabic_font(run_taawwudh, 28)

    p_bismillah = doc.add_paragraph()
    p_bismillah.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_bismillah = p_bismillah.add_run("بِسْمِ اللَّهِ الرَّحْمَـٰنِ الرَّحِيمِ")
    set_arabic_font(run_bismillah, 28)
    
    doc.add_paragraph("\n")
    
    # Dynamic Title Block
    p_title_block = doc.add_paragraph()
    p_title_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_main_title = p_title_block.add_run("Alimiyya Bayazan Workbook\n")
    run_main_title.bold = True
    run_main_title.font.size = Pt(28)
    run_main_title.font.color.rgb = RGBColor(33, 97, 140) 
    
    run_sub_title = p_title_block.add_run(f"Covering Surah {start_surah} to Surah {end_surah}\n\n")
    run_sub_title.font.size = Pt(18)
    run_sub_title.font.color.rgb = RGBColor(100, 100, 100) 
    
    # --- DATA SOURCE ATTRIBUTION ---
    p_attribution = doc.add_paragraph()
    p_attribution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_attr_label1 = p_attribution.add_run("Arabic Source Text: ")
    run_attr_label1.font.size = Pt(10)
    run_attr_label1.font.bold = True
    run_attr_label1.font.color.rgb = RGBColor(128, 128, 128)
    
    run_attr_url1 = p_attribution.add_run("https://tanzil.net/download/\n")
    run_attr_url1.font.size = Pt(10)
    run_attr_url1.font.color.rgb = RGBColor(5, 99, 193) 
    run_attr_url1.font.underline = True

    run_attr_label2 = p_attribution.add_run("Optimized for Font: KFGQPC Nastaleeq ")
    run_attr_label2.font.size = Pt(10)
    run_attr_label2.font.bold = True
    run_attr_label2.font.color.rgb = RGBColor(128, 128, 128)

    run_attr_url2 = p_attribution.add_run("(Download from King Fahd Complex)")
    run_attr_url2.font.size = Pt(10)
    run_attr_url2.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph("\n")
    
    # --- TABLE OF CONTENTS INJECTION ---
    p_index_title = doc.add_paragraph()
    p_index_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_index_title = p_index_title.add_run("Index / الفهرس")
    run_index_title.bold = True
    run_index_title.font.size = Pt(16)

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_note = p_note.add_run("📌 Quick Note: To calculate dynamic page numbers, right-click the grey text below and select 'Update Field'.")
    run_note.font.italic = True
    run_note.font.size = Pt(10)
    run_note.font.color.rgb = RGBColor(128, 128, 128)
    
    p_toc = doc.add_paragraph()
    p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_of_contents(p_toc)
    
    doc.add_page_break() 

    # --- READ SOURCE TEXT ---
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        lines = f.readlines()

    current_surah = None

    for line in lines:
        if not line.strip() or line.startswith('#'): continue
        
        parts = line.strip().split('|')
        if len(parts) < 3: continue
        
        s_num, a_num, ayah_text = int(parts[0]), int(parts[1]), parts[2]

        if s_num < start_surah: continue
        if s_num > end_surah: break 

        # --- BISMILLAH FIX ---
        if s_num != 1 and s_num != 9 and a_num == 1:
            temp_words = ayah_text.split()
            if len(temp_words) > 4:
                ayah_text = " ".join(temp_words[4:])

        # --- 1. SURAH HEADER BLOCK ---
        if current_surah != s_num:
            if current_surah is not None:
                doc.add_page_break()
            
            current_surah = s_num
            
            p_title = doc.add_paragraph(style='Heading 1')
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_title = p_title.add_run(f"سورة {s_num}")
            set_arabic_font(run_title, 36)
            run_title.bold = True
            run_title.font.color.rgb = RGBColor(25, 111, 61) 
            
            p_intro = doc.add_paragraph("Surah Introduction & Key Themes:")
            p_intro.bold = True
            p_intro.style.font.color.rgb = RGBColor(33, 97, 140) 
            doc.add_paragraph("\n\n\n") 

        # --- 2. AYAH HEADER ---
        doc.add_paragraph("_" * 80) 
        
        p_ayah = doc.add_paragraph()
        p_ayah.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_ayah.paragraph_format.line_spacing = Pt(40) 
        run_ayah = p_ayah.add_run(f"{ayah_text} ﴿{a_num}﴾")
        set_arabic_font(run_ayah, 26)

        # --- 3. WORD-BY-WORD TARKEEB TABLE ---
        words = ayah_text.split()
        
        table = doc.add_table(rows=len(words) + 1, cols=5)
        table.style = 'Table Grid'
        table.direction = WD_TABLE_DIRECTION.RTL
        
        widths = [Inches(1.5), Inches(1.0), Inches(2.5), Inches(2.0), Inches(3.0)]
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = width

        headers = ["الكلمة (Word)", "المادة (Root)", "التركيب (Role)", "المعنى (Meaning)", "ملاحظات (Notes)"]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            set_cell_background(cell, "E8F8F5") 
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.color.rgb = RGBColor(20, 90, 50)

        for row_idx, word in enumerate(words, start=1):
            row_cells = table.rows[row_idx].cells
            
            p_word = row_cells[0].paragraphs[0]
            p_word.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_word = p_word.add_run(word)
            set_arabic_font(run_word, 20)

        # --- 4. AYAH SPECIFIC TAFSEER ---
        doc.add_paragraph("\n") 
        p_notes = doc.add_paragraph("Ayah Context / General Tafseer:")
        p_notes.bold = True
        p_notes.style.font.color.rgb = RGBColor(33, 97, 140) 
        doc.add_paragraph("\n\n")

    doc.save(output_filepath)
    
    # Copy font file and create README
    copy_font_to_output()
    create_font_readme()
    
    print(f"Success! Final Indexed Workbook compiled and saved to {output_filepath}")

def copy_font_to_output():
    """Copy the standard mode font file to the generated directory."""
    output_dir = os.path.dirname(os.path.abspath(INPUT_FILE))
    output_dir = "generated"  # Use the generated directory
    
    font_source = PATHS["standard"]["font_file"]
    if os.path.exists(font_source):
        os.makedirs(output_dir, exist_ok=True)
        font_filename = os.path.basename(font_source)
        font_dest = os.path.join(output_dir, font_filename)
        
        # Only copy if it doesn't exist
        if not os.path.exists(font_dest):
            shutil.copy2(font_source, font_dest)
            print(f"📝 Font copied: {font_filename}")
    else:
        print(f"⚠️  Warning: Font file not found: {font_source}")

def create_font_readme():
    """Create a README file in the generated directory with font installation instructions."""
    output_dir = "generated"
    os.makedirs(output_dir, exist_ok=True)
    readme_path = os.path.join(output_dir, "FONT_INSTALLATION.md")
    
    # Only create if it doesn't exist
    if not os.path.exists(readme_path):
        font_name = PATHS["standard"]["font_name"]
        font_file = os.path.basename(PATHS["standard"]["font_file"])
        
        readme_content = f"""# Font Installation Required

## ⚠️ Important: Install Font Before Opening Documents

The generated Word documents in this directory use a specific Arabic font for proper text rendering. You **must** install the font before opening the documents to ensure correct display.

### Standard Mode Font
**Font Required:** `{font_name}`

### Installation Instructions

#### Windows:
1. Locate the font file: `{font_file}`
2. Right-click on the font file
3. Select "Install" or "Install for all users"
4. Restart Microsoft Word if it's already open

#### macOS:
1. Double-click the font file: `{font_file}`
2. Click "Install Font" in Font Book
3. Restart Microsoft Word if it's already open

#### Linux:
1. Copy the font file to `~/.fonts/` or `/usr/share/fonts/`
2. Run: `fc-cache -f -v`
3. Restart your document viewer

### Verification

After installing the font:
1. Open any generated `.docx` file
2. The Arabic text should display clearly with proper diacritics
3. If text appears as boxes or incorrect characters, the font is not properly installed

### Font Source

Download from: [King Fahd Glorious Qur'an Printing Complex](https://fonts.qurancomplex.gov.sa/)

### Troubleshooting

If the text still doesn't display correctly after installation:
1. Ensure you installed the correct font file
2. Restart your computer (not just Word)
3. Check that the font appears in your system's font list
4. Try opening the document in a different Word version

---

**Generated by Alimiyya Bayazan Generator - Standard Mode**
"""
        
        with open(readme_path, 'w', encoding='utf-8') as f:
            f.write(readme_content)
        print(f"📄 Created: FONT_INSTALLATION.md")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Bayazan Quran Word Docs")
    parser.add_argument("--start", type=int, required=True, help="Starting Surah Number (1-114)")
    parser.add_argument("--end", type=int, required=True, help="Ending Surah Number (1-114)")
    parser.add_argument("-o", "--output", type=str, help="Custom filename for the output Word document")
    
    args = parser.parse_args()
    
    if args.start < 1 or args.end > 114 or args.start > args.end:
        print("Error: Invalid Surah range. Must be between 1 and 114.")
    else:
        # Resolve the output directory
        output_dir = "generated"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # Resolve the filename
        if args.output:
            final_filename = args.output
            if not final_filename.endswith('.docx'):
                final_filename += '.docx'
        else:
            final_filename = f"Bayazan_Surahs_{args.start}_to_{args.end}.docx"
            
        # Combine directory and filename
        final_filepath = os.path.join(output_dir, final_filename)
            
        create_doc(args.start, args.end, final_filepath)
